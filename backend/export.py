import io
from typing import List
from docx import Document
from docx.shared import Pt

from backend.models import TerminRead


def generate_docx(termine: List[TerminRead]) -> io.BytesIO:
    doc = Document()

    # Title
    if termine:
        date_from = min(t.datum for t in termine).strftime("%d.%m.%Y")
        date_to = max(t.datum for t in termine).strftime("%d.%m.%Y")
        title = doc.add_paragraph(f"Ministrantenplan {date_from} - {date_to}")
        title.runs[0].bold = True
        title.runs[0].font.size = Pt(14)

    doc.add_paragraph("Gerne auch zum ministrieren kommen, wenn ihr nicht eingeteilt seid")

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Datum", "Uhrzeit", "Ministranten", "Ereignis"]):
        cell.text = text
        cell.paragraphs[0].runs[0].bold = True

    # Data rows
    for t in sorted(termine, key=lambda x: x.datum):
        row = table.add_row().cells

        # Datum + Wochentag
        datum_str = t.datum.strftime("%d.%m.%Y")
        row[0].text = f"{datum_str}\n{t.wochentag}"

        # Uhrzeit + Priester
        uhrzeit_str = t.uhrzeit
        if t.priester:
            uhrzeit_str += f"\n{t.priester}"
        row[1].text = uhrzeit_str

        # Ministranten
        if t.zuteilungen:
            row[2].text = " \u2013 ".join(z.name for z in t.zuteilungen)
        else:
            row[2].text = ""

        # Ereignis
        row[3].text = t.ereignis or ""

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
